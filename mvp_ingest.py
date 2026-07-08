"""
================================================================================
MVP 入库脚本 — 自包含的知识库入库流水线
================================================================================

零项目 import，所有逻辑内联。修改下方的 Config 配置，然后直接运行：

    python mvp_ingest.py

或者：

    uv run python mvp_ingest.py

----------------------
执行流程
----------------------

    文档加载 -> 文本清洗 -> 语义切块 -> ES 向量索引
    -> LLM 生成产品描述 -> LLM 生成客户问题 -> 批量检索 + 回答
    -> 去重归一化 -> QA 向量索引 -> 输出 JSON / Markdown / CSV

----------------------
输入格式
----------------------

    支持: .pdf  .docx  .txt  .md  .csv  .xlsx

    PDF 解析策略:
      1. 若配置了 mineru_server，优先使用远程 MinerU OCR
      2. MinerU 不可用或失败时，自动回退 pypdf
      3. 若 mineru_server 为空，直接使用 pypdf

----------------------
输出产物
----------------------

    {output_dir}/qa_list.json           结构化 QA 对（含证据来源）
    {output_dir}/qa_list.md             Markdown 格式 QA 文档
    {output_dir}/qa_list.csv            CSV 表格（可导入 Excel）
    {output_dir}/product_description.md  LLM 生成的产品介绍文档

    ES 索引:
      docs_{workspace_id}               文档向量索引（检索用）
      qa_{workspace_id}                 QA 向量索引（检索用）

----------------------
依赖
----------------------

    openai  elasticsearch[async]  pypdf  python-docx
    pandas  openpyxl  pypdfium2  mineru-vl-utils

    使用前确保 uv sync 已完成。
"""

from __future__ import annotations
import asyncio

import csv
import hashlib
import json
import re
import sys
import time
from concurrent.futures import ThreadPoolExecutor, as_completed
from contextlib import asynccontextmanager
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any


# ============================================================
# 配置 — 在这里修改参数，然后直接运行 python mvp_ingest.py
# ============================================================

@dataclass
class Config:
    # ---- 输入输出 ----
    input_dir: str = "./docs/"            # 文档目录，递归扫描子目录
    project: str = "省心说客服"            # 项目名称，决定 ES 索引命名空间
    output_dir: str = "mvp_output"        # 产物输出目录

    # ---- LLM ----
    llm_base_url: str = "http://183.147.142.111:30000/v1"
    llm_api_key: str = "EMPTY"
    llm_model: str = "glm-4.7"

    # ---- Embedding ----
    embedding_url: str = ""               # 留空则使用 llm_base_url
    embedding_model: str = "bge-m3"
    embedding_dim: int = 1024

    # ---- Elasticsearch ----
    # 方式1: Elastic Cloud — 填写 cloud_id + api_key（或 username/password）
    es_cloud_id: str = ""                       # Elastic Cloud 的 Cloud ID
    es_api_key: str = ""                        # API Key（优先于密码）
    # 方式2: 自建 / 其他云 — 填写完整 URL + 账号密码
    es_url: str = ""                            # 如 "https://xxx.es.aliyuncs.com:9200"
    es_username: str = "elastic"
    es_password: str = ""

    # ---- MinerU (PDF OCR，可选) ----
    mineru_server: str = ""               # 如 "http://x.x.x.x:63359"，留空则用 pypdf

    # ---- QA 生成参数 ----
    qa_limit: int = 12                    # 生成的核心问题数
    qa_generalization: int = 2            # 每个问题的变体数量
    chunk_size: int = 512                 # 文本切块大小
    chunk_overlap: int = 120              # 切块重叠字数

# ============================================================
# 1. 数据模型
# ============================================================

@dataclass(slots=True)
class DocumentPage:
    page_number: int | None
    text: str


@dataclass(slots=True)
class DocumentRecord:
    source_path: str
    file_name: str
    file_type: str
    title: str
    pages: list[DocumentPage]
    metadata: dict[str, str] = field(default_factory=dict)

    @property
    def raw_text(self) -> str:
        return "\n\n".join(p.text for p in self.pages if p.text.strip())


@dataclass(slots=True)
class TextChunk:
    id: str
    workspace_id: str
    file_name: str
    source_path: str
    content: str
    page_number: int | None
    section: str
    chunk_type: str
    metadata: dict[str, str] = field(default_factory=dict)


@dataclass(slots=True)
class RetrievedChunk:
    chunk: TextChunk
    score: float
    retrieval_method: str


@dataclass(slots=True)
class QAEvidence:
    chunk_id: str
    file_name: str
    excerpt: str
    page_number: int | None
    score: float

    def to_dict(self) -> dict[str, Any]:
        return {
            "chunk_id": self.chunk_id, "file_name": self.file_name,
            "excerpt": self.excerpt, "page_number": self.page_number, "score": self.score,
        }


@dataclass(slots=True)
class QAPair:
    question: str
    answer: str
    category: str
    risk_notes: list[str]
    evidence: list[QAEvidence] = field(default_factory=list)

    def to_dict(self) -> dict[str, Any]:
        return {
            "question": self.question, "answer": self.answer,
            "category": self.category, "risk_notes": self.risk_notes,
            "evidence": [e.to_dict() for e in self.evidence],
        }


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".text", ".md", ".csv", ".xlsx"}


# ============================================================
# 2. 工具函数
# ============================================================

def _short_hash(value: str) -> str:
    return hashlib.sha1(value.encode()).hexdigest()[:12]


def _utc_now() -> str:
    return time.strftime("%Y-%m-%dT%H:%M:%SZ", time.gmtime())


def _safe_excerpt(text: str, limit: int = 220) -> str:
    compact = re.sub(r"\s+", " ", (text or "").strip())
    return compact if len(compact) <= limit else compact[:limit] + "..."


# ============================================================
# 3. 文档加载
# ============================================================

def _load_pdf_pypdf(file_path: Path) -> list[DocumentPage]:
    from pypdf import PdfReader
    reader = PdfReader(str(file_path))
    return [DocumentPage(page_number=i, text=(p.extract_text() or ""))
            for i, p in enumerate(reader.pages, start=1)]


def _mineru_available(server_url: str) -> tuple[bool, str]:
    """检查 MinerU 运行时是否可用。"""
    if not server_url:
        return False, "未配置 --mineru-server"
    try:
        import importlib.util
        if importlib.util.find_spec("mineru_vl_utils") is None:
            return False, "缺少 mineru_vl_utils 包"
        if importlib.util.find_spec("pypdfium2") is None:
            return False, "缺少 pypdfium2 包"
    except Exception:
        return False, "依赖检查失败"
    return True, ""


def _load_pdf_mineru(file_path: Path, server_url: str, output_dir: Path) -> list[DocumentPage]:
    """用远程 MinerU 服务解析 PDF。失败抛异常，由上游回退到 pypdf。"""
    import pypdfium2 as pdfium
    from mineru_vl_utils import MinerUClient

    # 1. 渲染 PDF 为图片
    doc = pdfium.PdfDocument(str(file_path))
    images = []
    try:
        for i in range(len(doc)):
            page = doc.get_page(i)
            try:
                bitmap = page.render(scale=2.0)
                images.append(bitmap.to_pil().convert("RGB"))
            finally:
                page.close()
    finally:
        doc.close()

    if not images:
        raise RuntimeError("PDF 渲染为空")

    # 2. 逐页调用远程 MinerU
    client = MinerUClient(backend="http-client", server_url=server_url)
    start = time.monotonic()
    page_texts = []
    for idx, img in enumerate(images, start=1):
        if time.monotonic() - start > 300:
            raise RuntimeError(f"MinerU 超时 (>300s)")
        blocks = client.two_step_extract(img)
        text = "\n".join(str(b.get("content", "")).strip() for b in blocks if b.get("content"))
        page_texts.append(f"-- {idx} of {len(images)} --\n{text}")

    raw = "\n\n".join(page_texts)

    # 3. 按页标记拆页
    marker = re.compile(r"--\s*(\d+)\s+of\s+\d+\s*--")
    matches = list(marker.finditer(raw))
    if matches:
        pages = []
        for i, m in enumerate(matches):
            start_pos = m.end()
            end_pos = matches[i + 1].start() if i + 1 < len(matches) else len(raw)
            body = raw[start_pos:end_pos].strip()
            if body:
                pages.append(DocumentPage(page_number=int(m.group(1)), text=body))
        if pages:
            # 持久化到本地便于排查
            output_dir.mkdir(parents=True, exist_ok=True)
            (output_dir / "mineru_result.txt").write_text(raw, encoding="utf-8")
            return pages

    return [DocumentPage(page_number=1, text=raw)]


def _load_docx(file_path: Path) -> list[DocumentPage]:
    from docx import Document as DocxDocument
    doc = DocxDocument(file_path)  # ty:ignore[invalid-argument-type]
    segments = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                segments.append(" | ".join(cells))
    return [DocumentPage(page_number=1, text="\n".join(segments))]


def _load_csv(file_path: Path) -> list[DocumentPage]:
    import pandas as pd
    df = pd.read_csv(file_path).fillna("").astype(str)
    cols = " | ".join(df.columns.tolist())
    rows = [f"[{file_path.name}]\ncolumns: {cols}"]
    for rn, row in df.iterrows():
        cells = [f"{c}={row[c].strip()}" for c in df.columns if row[c].strip()]
        if cells:
            rows.append(f"row {rn + 1}: " + "; ".join(cells))  # ty:ignore[unsupported-operator]
    return [DocumentPage(page_number=1, text="\n".join(rows))]


def _load_xlsx(file_path: Path) -> list[DocumentPage]:
    import pandas as pd
    pages = []
    with pd.ExcelFile(file_path) as wb:
        for idx, sheet in enumerate(wb.sheet_names, start=1):
            df = wb.parse(sheet_name=sheet).fillna("").astype(str)  # ty:ignore[unresolved-attribute]
            cols = " | ".join(df.columns.tolist())
            rows = [f"[{file_path.name}:{sheet}]\ncolumns: {cols}"]
            for rn, row in df.iterrows():
                cells = [f"{c}={row[c].strip()}" for c in df.columns if row[c].strip()]
                if cells:
                    rows.append(f"row {rn + 1}: " + "; ".join(cells))  # ty:ignore[unsupported-operator]
            pages.append(DocumentPage(page_number=idx, text="\n".join(rows)))
    return pages


def _guess_title(pages: list[DocumentPage], fallback: str) -> str:
    for page in pages:
        for line in page.text.splitlines():
            if line.strip():
                return line.strip()[:80]
    return fallback


def load_docs(file_paths: list[Path], *, mineru_server: str = "",
              output_dir: Path | None = None) -> list[DocumentRecord]:
    documents: list[DocumentRecord] = []
    mineru_ok, mineru_msg = _mineru_available(mineru_server)
    if mineru_server and not mineru_ok:
        print(f"  [警告] MinerU 不可用: {mineru_msg}，PDF 将使用 pypdf")

    for fp in file_paths:
        ext = fp.suffix.lower()
        if ext == ".pdf":
            if mineru_ok:
                try:
                    out = (output_dir or Path(".")) / f"mineru_{_short_hash(fp.name)}"
                    pages = _load_pdf_mineru(fp, mineru_server, out)
                    print(f"  [MinerU] {fp.name}: {len(pages)} 页")
                except Exception as e:
                    print(f"  [MinerU] {fp.name} 失败，回退 pypdf: {e}")
                    pages = _load_pdf_pypdf(fp)
            else:
                pages = _load_pdf_pypdf(fp)
        elif ext == ".docx":
            pages = _load_docx(fp)
        elif ext in {".txt", ".text", ".md"}:
            pages = [DocumentPage(page_number=1, text=fp.read_text(encoding="utf-8", errors="ignore"))]
        elif ext == ".csv":
            pages = _load_csv(fp)
        elif ext == ".xlsx":
            pages = _load_xlsx(fp)
        else:
            continue
        title = _guess_title(pages, fp.stem)
        documents.append(DocumentRecord(
            source_path=str(fp), file_name=fp.name, file_type=ext.lstrip("."),
            title=title, pages=pages, metadata={"extension": ext.lstrip(".")},
        ))
    return documents


# ============================================================
# 4. 文本清洗
# ============================================================

_CONTROL_CHARS = re.compile(r"[\x01-\x08\x0b-\x1f\x7f-\x9f]")
_PAGE_MARKERS = re.compile(r"--\s*\d+\s+of\s+\d+\s*--")
_MULTI_SPACES = re.compile(r"[ \t]{2,}")
_MULTI_NEWLINES = re.compile(r"\n{3,}")


def _is_single_cjk(token: str) -> bool:
    return len(token) == 1 and bool(re.fullmatch(r"[一-鿿]", token))


def _normalize_line(line: str) -> str:
    stripped = line.strip()
    if not stripped:
        return ""
    tokens = [t for t in stripped.split() if t]
    if len(tokens) >= 3 and all(_is_single_cjk(t) for t in tokens):
        return "".join(tokens)
    return stripped


def _dedupe_lines(text: str) -> str:
    deduped, last = [], ""
    for line in text.splitlines():
        s = line.strip()
        if not s:
            deduped.append(""); last = ""; continue
        if s == last: continue
        deduped.append(s); last = s
    return "\n".join(deduped)


def clean_text(text: str) -> str:
    t = text.replace("\x00", "")  # 先去掉 null bytes
    t = t.replace("\r\n", "\n").replace("\r", "\n")
    t = t.replace("　", " ").replace("\xa0", " ")
    t = _CONTROL_CHARS.sub("", t)
    t = _PAGE_MARKERS.sub("\n", t)
    t = "\n".join(_normalize_line(l) for l in t.splitlines())
    t = _dedupe_lines(t)
    t = _MULTI_SPACES.sub(" ", t)
    t = _MULTI_NEWLINES.sub("\n\n", t)
    return t.strip()


def clean_docs(documents: list[DocumentRecord]) -> list[DocumentRecord]:
    return [
        DocumentRecord(
            source_path=d.source_path, file_name=d.file_name, file_type=d.file_type,
            title=clean_text(d.title),
            pages=[DocumentPage(page_number=p.page_number, text=clean_text(p.text)) for p in d.pages],
            metadata=d.metadata,
        )
        for d in documents
    ]


# ============================================================
# 5. 文档切块
# ============================================================

_HEADING_PATTERN = re.compile(r"^[一二三四五六七八九十\d]+[、.)）]?\s*|^第\s*\d+\s*[步章节项]")
_SENTENCE_BOUNDARY = re.compile(r"(?<=[。！？!?；;])")


def _looks_like_heading(text: str) -> bool:
    if len(text) > 42: return False
    if _HEADING_PATTERN.search(text): return True
    return text.endswith(("：", ":")) and len(text) <= 30


def _split_large(text: str, chunk_size: int, chunk_overlap: int) -> list[str]:
    sentences = [s.strip() for s in _SENTENCE_BOUNDARY.split(text) if s.strip()]
    if not sentences: return [text[:chunk_size]]
    pieces, buf = [], ""
    for s in sentences:
        if not buf: buf = s; continue
        cand = f"{buf}{s}"
        if len(cand) <= chunk_size: buf = cand; continue
        pieces.append(buf)
        overlap = buf[-chunk_overlap:] if chunk_overlap > 0 else ""
        buf = f"{overlap}{s}"
    if buf: pieces.append(buf)
    return [p.strip() for p in pieces if p.strip()]


def _build_chunk(doc: DocumentRecord, workspace_id: str, page_number: int | None,
                 section: str, content: str) -> TextChunk:
    ct = "table" if "columns:" in content or "row " in content else "paragraph"
    seed = f"{doc.source_path}:{page_number}:{section}:{content}"
    return TextChunk(
        id=f"chunk_{_short_hash(seed)}", workspace_id=workspace_id,
        file_name=doc.file_name, source_path=doc.source_path, content=content.strip(),
        page_number=page_number, section=section[:120], chunk_type=ct,
        metadata={"document_title": doc.title},
    )


def chunk_docs(documents: list[DocumentRecord], *, workspace_id: str,
               chunk_size: int, chunk_overlap: int) -> list[TextChunk]:
    chunks: list[TextChunk] = []
    for doc in documents:
        section = doc.title or doc.file_name
        for page in doc.pages:
            paragraphs = [p.strip() for p in re.split(r"\n{2,}", page.text) if p.strip()]
            buf = ""
            for para in paragraphs:
                p = re.sub(r"\s+", " ", para).strip()
                if not p: continue
                if _looks_like_heading(p): section = p[:120]; continue
                if len(p) > chunk_size:
                    if buf:
                        chunks.append(_build_chunk(doc, workspace_id, page.page_number, section, buf))
                        buf = ""
                    for piece in _split_large(p, chunk_size, chunk_overlap):
                        chunks.append(_build_chunk(doc, workspace_id, page.page_number, section, piece))
                    continue
                if not buf: buf = p; continue
                cand = f"{buf}\n{p}"
                if len(cand) <= chunk_size: buf = cand
                else:
                    chunks.append(_build_chunk(doc, workspace_id, page.page_number, section, buf))
                    buf = (f"{buf[-chunk_overlap:].strip()}\n{p}" if chunk_overlap > 0 else p)
            if buf:
                chunks.append(_build_chunk(doc, workspace_id, page.page_number, section, buf))
    # 去重
    seen_content, seen_ids, deduped = set(), set(), []
    for c in chunks:
        sig = c.content.strip()
        if sig in seen_content or c.id in seen_ids: continue
        seen_content.add(sig); seen_ids.add(c.id); deduped.append(c)
    return deduped


# ============================================================
# 6. LLM 客户端（OpenAI 兼容接口）
# ============================================================

_GLM_THINKING_OFF = {"thinking": {"type": "disabled"}}
_QWEN_THINKING_OFF = {"enable_thinking": False}

# 内置 prompt 模板
PROMPT_PRODUCT_DESC = """请阅读以下关于产品的多份文档内容，整合成一份全面、详细、结构清晰的产品介绍文档（Markdown格式）。
如果文档中包含多个产品，请尽量整合成系列介绍。

重点包含以下内容：
1. 产品名称（必须明确指出）
2. 产品概述与核心定位
3. 包含的单品及其规格
4. 核心卖点与主要功效
5. 核心成分及其作用
6. 适用人群与肤质
7. 使用方法与顺序
8. 注意事项与限制表达
9. 工厂、研发背景或资质信息（如有）

文档内容：
{context}

请直接输出 Markdown 内容，不要包含其他寒暄语。
确保第一行是类似 "# 产品名称" 的一级标题。"""

PROMPT_CUSTOMER_QUESTIONS = """请扮演一位对护肤品有较高关注度、追求效果但也关心安全的消费者（潜在客户）。
阅读以下产品介绍，针对该产品，提出 {count} 个你最关心的不同方面的问题。
请确保问题自然、口语化，就像真实客户咨询客服一样。
不要提供答案，只提供问题列表。

对于每个核心问题，请额外生成 {generalization_count} 个语义相似但表述不同的变体（泛化问题），用于丰富知识库。

输出格式要求：
请直接输出 JSON 对象格式，包含一个 "questions" 键，其值为列表。每个元素包含 "category"（分类）、"question"（核心问题）和 "variations"（变体列表）。

示例：
{{
  "questions": [
    {{
      "category": "功效",
      "question": "这个用了多久能看到淡纹效果？",
      "variations": ["大概要用几瓶才能去皱纹？", "见效快吗，多久能有变化？"]
    }},
    ...
  ]
}}

产品介绍内容：
{product_description}

请只输出 JSON，不要包含 Markdown 代码块标记（```json ... ```）。
确保生成的总核心问题数量为 {count} 个。"""

PROMPT_QA_EXTRACT = """你是该产品的专业客服人员，正在为客服团队生成可直接使用的标准回复。

请严格遵守以下要求：
1. 只能基于提供的产品资料回答，不能编造信息。
2. 回答必须像真实客服在和客户对话一样自然、亲切、专业，可以直接发给客户。
3. 绝对不要出现"根据现有证据""根据资料显示""根据产品设计原理""证据中提及""文件中说明"等学术性、分析性表述——客服不会这么说话。
4. 如果资料中有"以上内容仅为成分介绍，不代表产品功效"等合规限制，用自然的方式融入回答（如"需要提醒您的是""温馨提示"），而不是生硬引用。
5. 语气亲和但专业，像是熟悉产品的资深客服在耐心解答。
6. 不要输出项目符号或编号列表，只输出一段连贯的回答。
7. 回答要简洁实用，避免过度铺陈技术细节，突出客户最关心的要点。

产品：{product_name}
问题类别：{category}
问题：{question}

参考资料如下：
{context}

请输出客服标准回复："""

DEFAULT_QUESTIONS = [
    ("产品概览", "{name}是什么产品，包含哪些单品？"),
    ("核心卖点", "{name}的核心卖点和主要功效是什么？"),
    ("适用人群", "{name}适合哪些人群或肌肤状态？"),
    ("使用方法", "{name}的推荐使用顺序和早晚使用方法是什么？"),
    ("成分亮点", "{name}里提到的核心成分有哪些？"),
    ("规格信息", "{name}各单品的规格或容量分别是多少？"),
    ("功效表达", "{name}有哪些可以对外描述的功效信息？"),
    ("注意事项", "{name}在沟通时有哪些注意事项或限制表达？"),
    ("工厂资质", "{name}背后的工厂、研发或资质信息有哪些？"),
    ("客户问答", "如果客户问 {name} 是否适合熬夜暗沉和细纹人群，客服应该如何回答？"),
    ("客户问答", "如果客户问 {name} 和普通水乳精华有什么差别，客服可以如何介绍？"),
    ("客户问答", "如果客户问 {name} 为什么是一套五件，客服应该如何解释？"),
]


class LLMClient:
    """OpenAI 兼容的 LLM + Embedding 客户端。"""

    def __init__(self, *, base_url: str, api_key: str, model: str,
                 embedding_url: str = "", embedding_model: str = "",
                 embedding_dim: int = 1024):
        from openai import OpenAI
        self.model = model
        self.embedding_model = embedding_model
        self.embedding_dim = embedding_dim
        self._client = OpenAI(api_key=api_key, base_url=base_url, timeout=180.0, max_retries=1)
        emb_url = embedding_url or base_url
        self._emb_client = OpenAI(api_key=api_key, base_url=emb_url, timeout=180.0, max_retries=0)

    # -- embedding --

    def embed(self, texts: list[str]) -> list[list[float]]:
        if not texts or not self.embedding_model: return []
        embeddings: list[list[float]] = []
        batch_size = 64
        for i in range(0, len(texts), batch_size):
            batch = texts[i:i + batch_size]
            for attempt in range(3):
                try:
                    resp = self._emb_client.embeddings.create(model=self.embedding_model, input=batch)
                    embeddings.extend([d.embedding for d in resp.data])
                    break
                except Exception:
                    if attempt == 2:
                        print(f"  [警告] embedding 第{i // batch_size + 1}批失败，已重试3次")
                        embeddings.extend([[] for _ in batch])
                    else:
                        time.sleep(min(2.0 * (attempt + 1), 5.0))
        return embeddings

    # -- chat --

    def _thinking_off(self) -> dict:
        m = self.model.lower()
        return _GLM_THINKING_OFF if ("glm" in m or "kimi" in m) else _QWEN_THINKING_OFF

    def _extract_text(self, content: Any) -> str:
        if content is None: return ""
        if isinstance(content, str): return content.strip()
        if isinstance(content, list):
            parts = []
            for item in content:
                t = getattr(item, "text", None) if not isinstance(item, dict) else item.get("text")
                if isinstance(t, str) and t.strip(): parts.append(t.strip())
            return "\n".join(parts).strip()
        return str(content).strip()

    def chat(self, system: str, user: str, *, temperature: float = 0.3,
             max_tokens: int = 2048, json_mode: bool = False) -> str:
        kwargs: dict = dict(
            model=self.model, temperature=temperature, max_tokens=max_tokens,
            messages=[{"role": "system", "content": system}, {"role": "user", "content": user}],
            extra_body=self._thinking_off(),
        )
        if json_mode:
            kwargs["response_format"] = {"type": "json_object"}
        try:
            resp = self._client.chat.completions.create(**kwargs)
            return self._extract_text(resp.choices[0].message.content if resp.choices else "")
        except Exception as e:
            print(f"  [LLM错误] {e}")
            return ""

    # -- 产品描述 --

    def generate_product_description(self, context: str) -> str:
        return self.chat("你是一个专业的产品文档撰写助手。",
                         PROMPT_PRODUCT_DESC.format(context=context),
                         temperature=0.3)

    # -- 客户问题 --

    def _try_generate_questions(self, prompt: str) -> list[dict]:
        raw = self.chat("你是一个善于提问的潜在客户。请输出JSON格式。",
                        prompt, temperature=0.7, json_mode=True)
        if not raw: return []
        try:
            data = json.loads(raw.replace("```json", "").replace("```", "").strip())
            if isinstance(data, list): return [i for i in data if isinstance(i, dict)]
            if isinstance(data, dict) and "questions" in data:
                qs = data["questions"]
                return [i for i in qs if isinstance(i, dict)] if isinstance(qs, list) else []
        except json.JSONDecodeError:
            pass
        return []

    def generate_customer_questions(self, product_description: str,
                                    count: int, generalization_count: int) -> list[dict]:
        prompt = PROMPT_CUSTOMER_QUESTIONS.format(
            product_description=product_description,
            count=count, generalization_count=generalization_count,
        )
        best: list[dict] = []
        for attempt in range(3):
            qs = self._try_generate_questions(prompt)
            n = sum(1 for q in qs if isinstance(q.get("question"), str) and q["question"].strip())
            if n >= count: return qs
            if len(qs) > len(best): best = qs
            if attempt < 2: time.sleep(5)
        return best

    # -- 单条 QA 生成 --

    def generate_answer(self, *, product_name: str, question: str, category: str,
                        retrieved_chunks: list[RetrievedChunk]) -> tuple[str, list[str]]:
        if not retrieved_chunks:
            return ("未在当前资料中检索到足够证据，建议客服谨慎回复并人工复核。", ["无检索证据"])
        context = _format_context(retrieved_chunks)
        prompt = PROMPT_QA_EXTRACT.format(
            product_name=product_name, category=category, question=question, context=context,
        )
        answer = self.chat("你是一个谨慎的客服知识库助手，只能根据提供证据回答。",
                           prompt, temperature=0.2)
        if not answer:
            answer = _fallback_answer(question, retrieved_chunks)
        risk = _risk_notes(retrieved_chunks)
        return answer, risk

    # -- 批量 QA 生成 --

    def generate_answers_batch(self, *, product_name: str,
                               questions: list[tuple[str, str]],
                               retrieved_list: list[list[RetrievedChunk]],
                               ) -> list[tuple[str, list[str]]]:
        results: list = [("", [])] * len(questions)
        with ThreadPoolExecutor(max_workers=min(len(questions), 36)) as ex:
            futs = {
                ex.submit(self.generate_answer, product_name=product_name,
                          question=q[1], category=q[0],
                          retrieved_chunks=retrieved): i
                for i, (q, retrieved) in enumerate(zip(questions, retrieved_list))
            }
            for fut in as_completed(futs):
                idx = futs[fut]
                try:
                    results[idx] = fut.result()
                except Exception:
                    results[idx] = (
                        _fallback_answer(questions[idx][1], retrieved_list[idx]),
                        _risk_notes(retrieved_list[idx]),
                    )
        return results


def _fallback_answer(question: str, retrieved_chunks: list[RetrievedChunk]) -> str:
    combined = " ".join(c.chunk.content for c in retrieved_chunks[:2])
    sentences = [s.strip() for s in re.split(r"(?<=[。！？!?；;])", combined) if s.strip()]
    if not sentences:
        return _safe_excerpt(combined, 220)
    selected = []
    for s in sentences:
        if len("".join(selected)) + len(s) > 180:
            break
        selected.append(s)
    ans = "".join(selected).strip()
    prefix = "根据资料，"
    if question and ans and not ans.startswith(prefix):
        return prefix + ans
    return ans or _safe_excerpt(combined, 220)


def _risk_notes(retrieved_chunks: list[RetrievedChunk]) -> list[str]:
    combined = " ".join(c.chunk.content for c in retrieved_chunks[:3])
    notes: list[str] = []
    if "不代表产品功效" in combined:
        notes.append('资料中包含"成分介绍不代表产品功效"的限制表述，回复时避免夸大宣传。')
    if any(kw in combined for kw in ["改善", "抗皱", "修护"]):
        notes.append("资料中涉及功效类表述，回复时注意措辞合规。")
    return notes


def _format_context(chunks: list[RetrievedChunk]) -> str:
    rows = []
    for i, rc in enumerate(chunks, start=1):
        page = f"第{rc.chunk.page_number}页" if rc.chunk.page_number is not None else "页码未知"
        rows.append(f"[证据{i}] 文件={rc.chunk.file_name} | {page} | 分数={rc.score:.3f}\n{_safe_excerpt(rc.chunk.content, 400)}")
    return "\n\n".join(rows)


# ============================================================
# 7. Elasticsearch 存储
# ============================================================

DEFAULT_RRF_K = 60


@asynccontextmanager
async def _es_ctx(*, url: str = "", cloud_id: str = "", api_key: str = "",
                   username: str = "", password: str = ""):
    from elasticsearch import AsyncElasticsearch
    kwargs: dict = {"verify_certs": True, "request_timeout": 30}
    if cloud_id:
        kwargs["cloud_id"] = cloud_id.strip()
    elif url:
        kwargs["hosts"] = [url.strip().rstrip("/")]
    else:
        raise ValueError("必须提供 es_cloud_id 或 es_url")
    if api_key:
        kwargs["api_key"] = api_key.strip()
    elif username and password:
        kwargs["basic_auth"] = (username.strip(), password.strip())
    client = AsyncElasticsearch(**kwargs)
    try: yield client
    finally: await client.close()


def _docs_index(workspace_id: str) -> str:
    return f"docs_{workspace_id.lower()}"


def _qa_index(workspace_id: str) -> str:
    return f"qa_{workspace_id.lower()}"


def _run_async(coro):
    """在同步上下文中执行协程。"""
    try:
        asyncio.get_running_loop()
    except RuntimeError:
        return asyncio.new_event_loop().run_until_complete(coro)
    raise RuntimeError("不能在已有事件循环中同步调用 ES 方法")


class ESStore:
    """Elasticsearch 索引 + 检索。"""

    def __init__(self, workspace_id: str, *,
                 es_url: str = "", es_cloud_id: str = "",
                 es_user: str = "", es_pass: str = "", es_api_key: str = "",
                 embedding_dim: int = 1024):
        self.wid = workspace_id
        self.es_url = es_url
        self.es_cloud_id = es_cloud_id
        self.es_user = es_user
        self.es_pass = es_pass
        self.es_api_key = es_api_key
        self.emb_dim = embedding_dim

    # -- index documents --

    def _docs_mapping(self) -> dict:
        return {
            "properties": {
                "id": {"type": "keyword"}, "workspace_id": {"type": "keyword"},
                "file_name": {"type": "text"}, "source_path": {"type": "text"},
                "content": {"type": "text", "analyzer": "standard"},
                "page_number": {"type": "integer"}, "section": {"type": "text"},
                "chunk_type": {"type": "keyword"},
                "metadata": {"type": "object", "enabled": False},
                "embedding": {"type": "dense_vector", "dims": self.emb_dim,
                              "index": True, "similarity": "cosine"},
            }
        }

    def index_docs(self, chunks: list[TextChunk], embeddings: list[list[float]]):
        async def _run():
            async with _es_ctx(url=self.es_url, cloud_id=self.es_cloud_id,
                               api_key=self.es_api_key,
                               username=self.es_user, password=self.es_pass) as es:
                idx = _docs_index(self.wid)
                # 删旧建新
                await es.options(ignore_status=404).indices.delete(index=idx)
                await es.indices.create(index=idx, mappings=self._docs_mapping())
                # 批量写入
                for i in range(0, len(chunks), 100):
                    batch_c = chunks[i:i + 100]
                    batch_e = embeddings[i:i + 100]
                    ops = []
                    for c, e in zip(batch_c, batch_e):
                        ops.append({"index": {"_index": idx, "_id": c.id}})
                        doc = {"id": c.id, "workspace_id": c.wid, "file_name": c.file_name,  # ty:ignore[unresolved-attribute]
                               "source_path": c.source_path, "content": c.content,
                               "page_number": c.page_number, "section": c.section,
                               "chunk_type": c.chunk_type, "metadata": c.metadata}
                        if e: doc["embedding"] = e
                        ops.append(doc)
                    await es.bulk(operations=ops, refresh=True)
                print(f"  [索引] docs: {len(chunks)} 条写入完成")
        _run_async(_run())

    # -- search docs (vector + keyword, RRF fusion) --

    def search_docs(self, query: str, query_embedding: list[float],
                    top_k: int = 5, vector_k: int = 10, keyword_k: int = 20,
                    rrf_k: int = DEFAULT_RRF_K) -> list[RetrievedChunk]:
        """两路独立召回 + RRF 融合，与原始 QASearchService 逻辑一致。"""
        async def _run():
            async with _es_ctx(url=self.es_url, cloud_id=self.es_cloud_id,
                               api_key=self.es_api_key,
                               username=self.es_user, password=self.es_pass) as es:
                idx = _docs_index(self.wid)
                if not await es.indices.exists(index=idx):
                    return []

                # -- 向量路召回 --
                vector_hits_raw = []
                if query_embedding and vector_k > 0:
                    vector_body: dict = {
                        "size": vector_k,
                        "knn": {"field": "embedding", "query_vector": query_embedding,
                                "k": vector_k, "num_candidates": max(vector_k * 2, 50)},
                    }
                    try:
                        vr = await es.search(index=idx, body=vector_body)
                        vector_hits_raw = vr.get("hits", {}).get("hits", [])
                    except Exception as e:
                        print(f"  [ES] 向量召回失败: {e}")

                # -- 关键词路召回 --
                keyword_hits_raw = []
                if keyword_k > 0:
                    keyword_body = {
                        "size": keyword_k,
                        "query": {"match": {"content": {"query": query}}},
                    }
                    try:
                        kr = await es.search(index=idx, body=keyword_body)
                        keyword_hits_raw = kr.get("hits", {}).get("hits", [])
                    except Exception as e:
                        print(f"  [ES] 关键词召回失败: {e}")

                # -- ES hit -> RetrievedChunk --
                def _hit_to_chunk(hit: dict, method: str) -> RetrievedChunk:
                    s = hit.get("_source", {})
                    meta = s.get("metadata") or {}
                    return RetrievedChunk(
                        chunk=TextChunk(
                            id=str(s.get("id", "")),
                            workspace_id=str(s.get("workspace_id", "")),
                            file_name=str(s.get("file_name", "")),
                            source_path=str(s.get("source_path", "")),
                            content=str(s.get("content", "")),
                            page_number=s.get("page_number"),
                            section=str(s.get("section", "")),
                            chunk_type=str(s.get("chunk_type", "paragraph")),
                            metadata={str(k): str(v) for k, v in meta.items()},
                        ),
                        score=float(hit.get("_score", 0)),
                        retrieval_method=method,
                    )

                # -- RRF 融合 --
                def _rrf(hits: list, route_name: str) -> dict[str, dict]:
                    scored: dict[str, dict] = {}
                    for rank, h in enumerate(hits, start=1):
                        cid = str(h.get("_source", {}).get("id", ""))
                        if not cid:
                            continue
                        if cid not in scored:
                            scored[cid] = {"item": h, "rrf": 0.0, "raw": 0.0}
                        scored[cid]["rrf"] += 1.0 / (rrf_k + rank)
                        score = float(h.get("_score", 0))
                        if score > scored[cid]["raw"]:
                            scored[cid]["raw"] = score
                            scored[cid]["item"] = h
                    return scored

                vec_map = _rrf(vector_hits_raw[:vector_k], "vector")
                kw_map = _rrf(keyword_hits_raw[:keyword_k], "keyword")

                all_keys = set(vec_map.keys()) | set(kw_map.keys())
                merged = []
                for cid in all_keys:
                    vs = vec_map.get(cid, {})
                    ks = kw_map.get(cid, {})
                    rrf_score = vs.get("rrf", 0.0) + ks.get("rrf", 0.0)
                    best_hit = vs.get("item") or ks.get("item")
                    if best_hit is None:
                        continue
                    route = "rrf"
                    if cid in vec_map and cid not in kw_map:
                        route = "elasticsearch_vector"
                    elif cid in kw_map and cid not in vec_map:
                        route = "elasticsearch_keyword"
                    merged.append((rrf_score, best_hit, route))

                merged.sort(key=lambda x: x[0], reverse=True)
                return [_hit_to_chunk(h, route) for _, h, route in merged[:top_k]]

        return _run_async(_run())

    # -- QA embedding 映射 --

    def _qa_mapping(self) -> dict:
        return {
            "properties": {
                "id": {"type": "keyword"}, "workspace_id": {"type": "keyword"},
                "question": {"type": "text", "analyzer": "standard"},
                "answer": {"type": "text", "analyzer": "standard"},
                "category": {"type": "keyword"}, "risk_notes": {"type": "text"},
                "evidence": {"type": "object", "enabled": False},
                "metadata": {"type": "object", "enabled": False},
                "qa_embedding": {"type": "dense_vector", "dims": self.emb_dim,
                                 "index": True, "similarity": "cosine"},
            }
        }

    def index_qa(self, qa_pairs: list[QAPair], embeddings: list[list[float]]):
        async def _run():
            async with _es_ctx(url=self.es_url, cloud_id=self.es_cloud_id,
                               api_key=self.es_api_key,
                               username=self.es_user, password=self.es_pass) as es:
                idx = _qa_index(self.wid)
                await es.options(ignore_status=404).indices.delete(index=idx)
                await es.indices.create(index=idx, mappings=self._qa_mapping())
                ops = []
                for qa, emb in zip(qa_pairs, embeddings):
                    doc_id = f"qa_{hashlib.md5(qa.question.encode()).hexdigest()[:16]}"
                    doc = {"id": doc_id, "workspace_id": self.wid, "question": qa.question,
                           "answer": qa.answer, "category": qa.category,
                           "risk_notes": " | ".join(qa.risk_notes),
                           "evidence": json.dumps([e.to_dict() for e in qa.evidence], ensure_ascii=False)}
                    if emb: doc["qa_embedding"] = emb
                    ops.append({"index": {"_index": idx, "_id": doc_id}})
                    ops.append(doc)
                if ops:
                    await es.bulk(operations=ops, refresh=True)
                print(f"  [索引] qa: {len(qa_pairs)} 条写入完成")
        _run_async(_run())


# ============================================================
# 8. QA 去重
# ============================================================

def _normalize_question(q: str) -> str:
    n = re.sub(r"\s+", "", q)
    n = re.sub(r"[？?！!。,，：:；;、]", "", n)
    return n.lower()


def dedupe_qa(qa_pairs: list[QAPair]) -> list[QAPair]:
    merged: dict[str, QAPair] = {}
    for pair in qa_pairs:
        key = _normalize_question(pair.question)
        if key not in merged:
            merged[key] = pair; continue
        existing = merged[key]
        # 合并 evidence
        ev_dict: dict[tuple[str, str], QAEvidence] = {}
        for e in existing.evidence + pair.evidence:
            k = (e.chunk_id, e.file_name)
            if k not in ev_dict or e.score > ev_dict[k].score:
                ev_dict[k] = e
        risk = sorted(set(existing.risk_notes + pair.risk_notes))
        ans = existing.answer if len(existing.answer) >= len(pair.answer) else pair.answer
        merged[key] = QAPair(question=existing.question, answer=ans,
                             category=existing.category, risk_notes=risk,
                             evidence=list(ev_dict.values()))
    return list(merged.values())


# ============================================================
# 9. QA 批量生成
# ============================================================

def generate_qa_pairs(*, product_name: str, llm: LLMClient, es: ESStore,
                      questions: list[tuple[str, str]], top_k: int = 5,
                      batch_size: int = 36) -> list[QAPair]:
    qa_pairs: list[QAPair] = []
    total = len(questions)

    for bstart in range(0, total, batch_size):
        bend = min(bstart + batch_size, total)
        batch = questions[bstart:bend]
        print(f"  [QA] 批次 [{bstart + 1}-{bend}]/{total}")

        # 逐问题检索
        retrieved_list: list[list[RetrievedChunk]] = []
        for idx, (category, question) in enumerate(batch, start=bstart + 1):
            print(f"  [QA] [{idx}] {question}")
            q_emb = llm.embed([question])
            q_vec = q_emb[0] if q_emb else []
            retrieved = es.search_docs(question, q_vec, top_k=top_k)
            retrieved_list.append(retrieved)

        # 批量回答
        results = llm.generate_answers_batch(
            product_name=product_name,
            questions=batch,
            retrieved_list=retrieved_list,
        )

        # 组装 QAPair，附带证据片段
        for (category, question), retrieved, (answer, risk_notes) in zip(
            batch, retrieved_list, results
        ):
            evidence = [
                QAEvidence(
                    chunk_id=rc.chunk.id,
                    file_name=rc.chunk.file_name,
                    excerpt=_safe_excerpt(rc.chunk.content),
                    page_number=rc.chunk.page_number,
                    score=rc.score,
                )
                for rc in retrieved[:top_k]
            ]
            qa_pairs.append(QAPair(
                question=question, answer=answer, category=category,
                risk_notes=risk_notes, evidence=evidence,
            ))

    return qa_pairs


# ============================================================
# 10. QA 输出格式化
# ============================================================

def qa_pairs_to_markdown(product_name: str, qa_pairs: list[QAPair]) -> str:
    lines = [f"# {product_name} QA 列表", ""]
    for i, pair in enumerate(qa_pairs, start=1):
        lines.append(f"## {i}. {pair.question}")
        lines.append(pair.answer)
        if pair.risk_notes:
            lines.append(""); lines.append("风险提示：")
            for n in pair.risk_notes: lines.append(f"- {n}")
        if pair.evidence:
            lines.append(""); lines.append("证据来源：")
            for e in pair.evidence:
                page = f"第{e.page_number}页" if e.page_number is not None else "页码未知"
                lines.append(f"- {e.file_name} / {page} / score={e.score:.3f} / {e.excerpt}")
        lines.append("")
    return "\n".join(lines).strip() + "\n"


def qa_pairs_to_rows(qa_pairs: list[QAPair]) -> list[dict[str, str]]:
    return [{
        "question": p.question, "answer": p.answer, "category": p.category,
        "risk_notes": " | ".join(p.risk_notes),
        "evidence_files": " | ".join({e.file_name for e in p.evidence}),
    } for p in qa_pairs]


# ============================================================
# 11. CLI + main
# ============================================================

def main():
    cfg = Config()

    input_dir = Path(cfg.input_dir).resolve()
    if not input_dir.is_dir():
        print(f"错误: input_dir 路径不存在: {input_dir}")
        raise SystemExit(1)

    print(f"[输入]   扫描 {input_dir}")
    file_paths = [fp for fp in input_dir.rglob("*")
                  if fp.is_file() and fp.suffix.lower() in SUPPORTED_EXTENSIONS]
    file_paths.sort(key=lambda p: p.name)
    if not file_paths:
        print("错误: 未找到受支持的文档文件")
        raise SystemExit(1)
    for fp in file_paths:
        print(f"         - {fp.name}")

    output_dir = Path(cfg.output_dir).resolve()
    output_dir.mkdir(parents=True, exist_ok=True)
    print(f"[输出]   {output_dir}")

    workspace_id = hashlib.md5(cfg.project.encode()).hexdigest()[:16]
    print(f"[ID]     {workspace_id}")

    # ── 初始化 LLM + ES ────────────────────────────────
    embedding_url = cfg.embedding_url or cfg.llm_base_url
    print(f"[LLM]    {cfg.llm_model} @ {cfg.llm_base_url}")
    print(f"[Embed]  {cfg.embedding_model} @ {embedding_url}")
    print(f"[ES]     {cfg.es_url}")

    llm = LLMClient(
        base_url=cfg.llm_base_url, api_key=cfg.llm_api_key, model=cfg.llm_model,
        embedding_url=embedding_url, embedding_model=cfg.embedding_model,
        embedding_dim=cfg.embedding_dim,
    )
    es = ESStore(workspace_id=workspace_id, es_url=cfg.es_url,
                 es_cloud_id=cfg.es_cloud_id,
                 es_user=cfg.es_username, es_pass=cfg.es_password,
                 es_api_key=cfg.es_api_key,
                 embedding_dim=cfg.embedding_dim)

    # ═══════════ 功能1: 加载 + 清洗 ═══════════
    print("\n── 功能1: 文档加载 + 清洗 ──")
    print(f"  MinerU: {'禁用' if not cfg.mineru_server else cfg.mineru_server}")
    print("  加载文档...")
    docs = load_docs(file_paths, mineru_server=cfg.mineru_server, output_dir=output_dir)
    if not docs:
        print("错误: 没有可处理的文档"); raise SystemExit(1)
    print(f"  已加载 {len(docs)} 个文档")

    print("  清洗文档...")
    docs = clean_docs(docs)

    # ═══════════ 生成产品描述 ═══════════
    print("\n── 产品描述生成 ──")
    all_text = "\n\n".join(d.raw_text for d in docs)
    desc = llm.generate_product_description(all_text[:100000])

    product_name = "未命名产品"
    for line in (desc or "").split("\n"):
        if line.strip().startswith("# "):
            product_name = line.strip().replace("# ", "").strip(); break
    print(f"  产品名: {product_name}")

    # 产品描述加入检索语料
    docs.append(DocumentRecord(
        source_path="generated/product_description.md", file_name="product_description.md",
        file_type="markdown", title="产品介绍文档",
        pages=[DocumentPage(page_number=1, text=desc)], metadata={"type": "generated"},
    ))

    # ═══════════ 切块 ═══════════
    print(f"\n── 文档切块 (size={cfg.chunk_size}, overlap={cfg.chunk_overlap}) ──")
    chunks = chunk_docs(docs, workspace_id=workspace_id,
                        chunk_size=cfg.chunk_size, chunk_overlap=cfg.chunk_overlap)
    print(f"  共 {len(chunks)} 个 chunk")
    if not chunks:
        print("错误: 切块结果为空"); raise SystemExit(1)

    # ═══════════ 功能3阶段A: docs 写入 ES ═══════════
    print("\n── 功能3: docs 写入 ES ──")
    chunk_texts = [c.content for c in chunks]
    print(f"  生成 embedding ({len(chunks)} 条)...")
    chunk_embs = llm.embed(chunk_texts)
    es.index_docs(chunks, chunk_embs)

    # ═══════════ 功能2: 生成 QA ═══════════
    print("\n── 功能2: QA 生成 ──")
    print("  生成客户问题...")
    questions_data = llm.generate_customer_questions(
        desc, count=cfg.qa_limit, generalization_count=cfg.qa_generalization,
    )
    questions_list = []
    for item in questions_data:
        cat = item.get("category", "通用")
        q = item.get("question", "")
        if q: questions_list.append((cat, q))
        for v in item.get("variations", []):
            if v: questions_list.append((cat, v))
    if not questions_list:
        questions_list = [(c, q.format(name=product_name)) for c, q in DEFAULT_QUESTIONS]
    print(f"  共 {len(questions_list)} 个问题（含变体）")

    print("  批量检索 + 生成答案...")
    qa_pairs = generate_qa_pairs(
        product_name=product_name, llm=llm, es=es,
        questions=questions_list, batch_size=cfg.qa_limit + cfg.qa_limit * cfg.qa_generalization,
    )

    print("  去重...")
    qa_pairs = dedupe_qa(qa_pairs)
    print(f"  去重后 {len(qa_pairs)} 条 QA")

    # ═══════════ 功能3阶段B: QA 写入 ES ═══════════
    print("\n── 功能3: QA 写入 ES ──")
    qa_texts = [f"{qa.question} {qa.answer}" for qa in qa_pairs]
    print(f"  生成 QA embedding ({len(qa_pairs)} 条)...")
    qa_embs = llm.embed(qa_texts)
    es.index_qa(qa_pairs, qa_embs)

    # ═══════════ 输出 ═══════════
    print("\n── 产物输出 ──")
    output_paths = {}

    jp = output_dir / "qa_list.json"
    jp.write_text(json.dumps({
        "product_name": product_name,
        "qa_pairs": [p.to_dict() for p in qa_pairs],
    }, ensure_ascii=False, indent=2), encoding="utf-8")
    output_paths["json"] = str(jp)

    mp = output_dir / "qa_list.md"
    mp.write_text(qa_pairs_to_markdown(product_name, qa_pairs), encoding="utf-8")
    output_paths["markdown"] = str(mp)

    cp = output_dir / "qa_list.csv"
    with cp.open("w", encoding="utf-8", newline="") as f:
        w = csv.DictWriter(f, fieldnames=["question", "answer", "category", "risk_notes", "evidence_files"])
        w.writeheader(); w.writerows(qa_pairs_to_rows(qa_pairs))
    output_paths["csv"] = str(cp)

    dp = output_dir / "product_description.md"
    dp.write_text(desc, encoding="utf-8")
    output_paths["product_description"] = str(dp)

    for label, path in sorted(output_paths.items()):
        print(f"  {label}: {path}")

    print(f"\n{'=' * 60}")
    print("[完成]")
    print(f"  产品名:   {product_name}")
    print(f"  Chunk 数: {len(chunks)}")
    print(f"  QA 对数:  {len(qa_pairs)}")
    print(f"  输出目录: {output_dir}")


if __name__ == "__main__":
    main()
