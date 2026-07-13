from __future__ import annotations

import re
import sys
from pathlib import Path
from typing import Callable

from simbrain.models import DocumentPage, DocumentRecord
from simbrain.utils import _short_hash


def _load_pdf_pypdf(file_path: Path) -> list[DocumentPage]:
    from pypdf import PdfReader

    reader = PdfReader(str(file_path))
    return [
        DocumentPage(page_number=i, text=(p.extract_text() or ""))
        for i, p in enumerate(reader.pages, start=1)
    ]


def _mineru_available(api_token: str) -> tuple[bool, str]:
    """检查 MinerU 云 API 是否可用。"""
    if not api_token:
        return False, "未配置 mineru_api_token"
    try:
        import importlib.util

        if importlib.util.find_spec("mineru") is None:
            return False, "缺少 mineru-open-sdk 包，请执行: uv add mineru-open-sdk"
    except Exception:
        return False, "依赖检查失败"
    return True, ""


def _load_pdf_mineru(file_path: Path, api_token: str, output_dir: Path) -> list[DocumentPage]:
    """用 MinerU 云 API 解析 PDF。失败抛异常，由上游回退到 pypdf。"""
    from mineru import MinerU

    client = MinerU(token=api_token)
    try:
        result = client.extract(str(file_path), model="vlm")
    except Exception:
        client.close()
        raise
    client.close()

    if result.state != "done" or not result.markdown:
        raise RuntimeError(f"MinerU 云解析失败: {result.error or '无内容返回'}")

    markdown = result.markdown.strip()

    output_dir.mkdir(parents=True, exist_ok=True)
    (output_dir / "mineru_result.md").write_text(markdown, encoding="utf-8")

    return [DocumentPage(page_number=1, text=markdown)]


def _load_docx(file_path: Path) -> list[DocumentPage]:
    from docx import Document as DocxDocument

    doc = DocxDocument(file_path)  # ty:ignore[invalid-argument-type]
    segments = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                segments.append(" | ".join(cells))
    return [DocumentPage(page_number=1, text="\n".join(segments))]


def _load_csv(file_path: Path) -> list[DocumentPage]:
    import pandas as pd

    df = pd.read_csv(file_path).fillna("").astype(str)
    cols = " | ".join(df.columns.tolist())
    rows = [f"[{file_path.name}]\ncolumns: {cols}"]
    for rn, row in df.iterrows():
        cells = [f"{c}={row[c].strip()}" for c in df.columns if row[c].strip()]
        if cells:
            rows.append(f"row {rn + 1}: " + "; ".join(cells))  # ty:ignore[unsupported-operator]
    return [DocumentPage(page_number=1, text="\n".join(rows))]


def _load_xlsx(file_path: Path) -> list[DocumentPage]:
    import pandas as pd

    pages = []
    with pd.ExcelFile(file_path) as wb:
        for idx, sheet in enumerate(wb.sheet_names, start=1):
            df = wb.parse(sheet_name=sheet).fillna("").astype(str)  # ty:ignore[unresolved-attribute]
            cols = " | ".join(df.columns.tolist())
            rows = [f"[{file_path.name}:{sheet}]\ncolumns: {cols}"]
            for rn, row in df.iterrows():
                cells = [f"{c}={row[c].strip()}" for c in df.columns if row[c].strip()]
                if cells:
                    rows.append(f"row {rn + 1}: " + "; ".join(cells))  # ty:ignore[unsupported-operator]
            pages.append(DocumentPage(page_number=idx, text="\n".join(rows)))
    return pages


def _guess_title(pages: list[DocumentPage], fallback: str) -> str:
    for page in pages:
        for line in page.text.splitlines():
            if line.strip():
                return line.strip()[:80]
    return fallback


def load_docs(
    file_paths: list[Path],
    *,
    mineru_api_token: str = "",
    output_dir: Path | None = None,
    source_root: Path | None = None,
    progress_callback: Callable[[int, int, str], None] | None = None,
) -> list[DocumentRecord]:
    documents: list[DocumentRecord] = []
    mineru_ok, mineru_msg = _mineru_available(mineru_api_token)
    if mineru_api_token and not mineru_ok:
        print(f"  [警告] MinerU 不可用: {mineru_msg}，PDF 将使用 pypdf", file=sys.stderr)

    total = len(file_paths)
    for position, fp in enumerate(file_paths, start=1):
        if progress_callback:
            progress_callback(position - 1, total, fp.name)
        ext = fp.suffix.lower()
        parser = ext.lstrip(".")
        mineru_artifact = ""
        if ext == ".pdf":
            if mineru_ok:
                try:
                    safe_stem = re.sub(r"[^\w.-]+", "_", fp.stem, flags=re.UNICODE).strip("._") or "document"
                    artifact_rel = Path("mineru") / f"{safe_stem}_{_short_hash(fp.name.casefold())}" / "mineru_result.md"
                    out = (output_dir or Path(".")) / artifact_rel.parent
                    pages = _load_pdf_mineru(fp, mineru_api_token, out)
                    parser = "mineru"
                    mineru_artifact = artifact_rel.as_posix()
                    print(f"  [MinerU] {fp.name}: {len(pages)} 页", file=sys.stderr)
                except Exception as e:
                    print(f"  [MinerU] {fp.name} 失败，回退 pypdf: {e}", file=sys.stderr)
                    pages = _load_pdf_pypdf(fp)
                    parser = "pypdf"
            else:
                pages = _load_pdf_pypdf(fp)
                parser = "pypdf"
        elif ext == ".docx":
            pages = _load_docx(fp)
        elif ext in {".txt", ".text", ".md"}:
            pages = [DocumentPage(page_number=1, text=fp.read_text(encoding="utf-8", errors="ignore"))]
        elif ext == ".csv":
            pages = _load_csv(fp)
        elif ext == ".xlsx":
            pages = _load_xlsx(fp)
        else:
            continue
        title = _guess_title(pages, fp.stem)
        source_path = str(fp.relative_to(source_root)) if source_root else str(fp)
        documents.append(
            DocumentRecord(
                source_path=source_path,
                file_name=fp.name,
                file_type=ext.lstrip("."),
                title=title,
                pages=pages,
                metadata={
                    "extension": ext.lstrip("."),
                    "parser": parser,
                    "mineru_artifact": mineru_artifact,
                },
            )
        )
        if progress_callback:
            progress_callback(position, total, fp.name)
    return documents
